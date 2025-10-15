from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from datetime import datetime
import io

class ExportManager:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def export_to_pdf(self, transcript, summary=None, topics=None):
        """Export meeting data to PDF format"""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        story.append(Paragraph("🎤 Meeting Summary Report", title_style))
        story.append(Spacer(1, 20))
        
        # Meeting info
        info_style = self.styles['Normal']
        story.append(Paragraph(f"<b>Generated on:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", info_style))
        story.append(Spacer(1, 20))
        
        # Executive Summary
        if summary:
            story.append(Paragraph("📋 Executive Summary", self.styles['Heading2']))
            story.append(Paragraph(summary.get('summary', 'Summary not available'), info_style))
            story.append(Spacer(1, 20))
            
            # Action Items
            if summary.get('action_items'):
                story.append(Paragraph("✅ Action Items", self.styles['Heading2']))
                for item in summary['action_items']:
                    story.append(Paragraph(f"• {item}", info_style))
                story.append(Spacer(1, 20))
            
            # Key Decisions
            if summary.get('key_decisions'):
                story.append(Paragraph("🔑 Key Decisions", self.styles['Heading2']))
                for decision in summary['key_decisions']:
                    story.append(Paragraph(f"• {decision}", info_style))
                story.append(Spacer(1, 20))
        
        # Topics
        if topics:
            story.append(Paragraph("🎯 Meeting Topics", self.styles['Heading2']))
            for topic in topics:
                story.append(Paragraph(f"<b>{topic['timestamp']} - {topic['title']}</b>", self.styles['Heading3']))
                story.append(Paragraph(f"Duration: {topic['duration']}", info_style))
                story.append(Paragraph(topic['summary'], info_style))
                story.append(Spacer(1, 15))
        
        # Full Transcript
        story.append(PageBreak())
        story.append(Paragraph("📝 Full Meeting Transcript", self.styles['Heading2']))
        story.append(Spacer(1, 20))
        
        # Split transcript into chunks to avoid memory issues
        transcript_chunks = [transcript[i:i+2000] for i in range(0, len(transcript), 2000)]
        for chunk in transcript_chunks:
            story.append(Paragraph(chunk, info_style))
        
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
    
    def export_to_docx(self, transcript, summary=None, topics=None):
        """Export meeting data to Word document format"""
        
        doc = Document()
        
        # Title
        title = doc.add_heading('🎤 Meeting Summary Report', 0)
        title.alignment = 1  # Center alignment
        
        # Meeting info
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph()
        
        # Executive Summary
        if summary:
            doc.add_heading('📋 Executive Summary', level=1)
            doc.add_paragraph(summary.get('summary', 'Summary not available'))
            
            # Action Items
            if summary.get('action_items'):
                doc.add_heading('✅ Action Items', level=1)
                for item in summary['action_items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)
            
            # Key Decisions
            if summary.get('key_decisions'):
                doc.add_heading('🔑 Key Decisions', level=1)
                for decision in summary['key_decisions']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(decision)
        
        # Topics
        if topics:
            doc.add_heading('🎯 Meeting Topics', level=1)
            for topic in topics:
                doc.add_heading(f"{topic['timestamp']} - {topic['title']}", level=2)
                doc.add_paragraph(f"Duration: {topic['duration']}")
                doc.add_paragraph(topic['summary'])
                doc.add_paragraph()
        
        # Full Transcript
        doc.add_page_break()
        doc.add_heading('📝 Full Meeting Transcript', level=1)
        doc.add_paragraph(transcript)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_data = buffer.getvalue()
        buffer.close()
        
        return docx_data
    
    def prepare_email(self, transcript, summary=None, topics=None):
        """Prepare email content for meeting summary"""
        
        email_content = f"""Subject: Meeting Summary - {datetime.now().strftime('%B %d, %Y')}

Dear Team,

Please find below the summary of our meeting held on {datetime.now().strftime('%B %d, %Y')}.

"""
        
        # Executive Summary
        if summary:
            email_content += """📋 EXECUTIVE SUMMARY
====================
"""
            email_content += f"{summary.get('summary', 'Summary not available')}\n\n"
            
            # Action Items
            if summary.get('action_items'):
                email_content += """✅ ACTION ITEMS
===============
"""
                for i, item in enumerate(summary['action_items'], 1):
                    email_content += f"{i}. {item}\n"
                email_content += "\n"
            
            # Key Decisions
            if summary.get('key_decisions'):
                email_content += """🔑 KEY DECISIONS
================
"""
                for i, decision in enumerate(summary['key_decisions'], 1):
                    email_content += f"{i}. {decision}\n"
                email_content += "\n"
        
        # Topics Overview
        if topics:
            email_content += """🎯 MEETING AGENDA & TOPICS
==========================
"""
            for topic in topics:
                email_content += f"{topic['timestamp']} - {topic['title']} ({topic['duration']})\n"
                email_content += f"   → {topic['summary']}\n\n"
        
        email_content += """📝 FULL TRANSCRIPT
==================
The complete meeting transcript is attached as a separate document.

Best regards,
AI Meeting Summarizer

---
This summary was automatically generated using AI technology.
Please review for accuracy and completeness.
"""
        
        return email_content
    
    def export_comprehensive_pdf(self, export_data: dict, language: str = "English"):
        """Export comprehensive meeting data including all new features to PDF"""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1
        )
        
        story.append(Paragraph(f"🎤 Complete Meeting Report ({language})", title_style))
        story.append(Spacer(1, 20))
        
        # Meeting metadata
        info_style = self.styles['Normal']
        story.append(Paragraph(f"<b>Generated on:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", info_style))
        story.append(Paragraph(f"<b>Language:</b> {language}", info_style))
        story.append(Spacer(1, 20))
        
        # Risk Analysis Section
        if export_data.get('risks'):
            story.append(Paragraph("🚨 Risk Analysis", self.styles['Heading2']))
            risks = export_data['risks']
            
            # Risk summary
            total_risks = risks.get('total_risks', 0)
            urgency_score = risks.get('urgency_score', 0)
            
            story.append(Paragraph(f"<b>Total Risk Items:</b> {total_risks}", info_style))
            story.append(Paragraph(f"<b>Urgency Score:</b> {urgency_score}/100", info_style))
            story.append(Spacer(1, 10))
            
            # Individual risk categories
            risk_categories = [
                ('Deadlines', risks.get('deadlines', [])),
                ('Budget Risks', risks.get('budget_risks', [])),
                ('Legal Concerns', risks.get('legal_concerns', [])),
                ('Customer Issues', risks.get('customer_issues', []))
            ]
            
            for category, items in risk_categories:
                if items:
                    story.append(Paragraph(f"<b>{category}:</b>", self.styles['Heading3']))
                    for item in items:
                        story.append(Paragraph(f"• {item}", info_style))
                    story.append(Spacer(1, 10))
            
            story.append(Spacer(1, 20))
        
        # Speaker Information
        if export_data.get('speakers'):
            story.append(Paragraph("👥 Meeting Participants", self.styles['Heading2']))
            speakers = export_data['speakers']
            for old_name, new_name in speakers.items():
                if old_name != new_name:
                    story.append(Paragraph(f"• {new_name} ({old_name})", info_style))
                else:
                    story.append(Paragraph(f"• {new_name}", info_style))
            story.append(Spacer(1, 20))
        
        # Executive Summary
        summary = export_data.get('summary')
        if summary:
            story.append(Paragraph("📋 Executive Summary", self.styles['Heading2']))
            story.append(Paragraph(summary.get('summary', 'Summary not available'), info_style))
            story.append(Spacer(1, 20))
            
            # Action Items
            if summary.get('action_items'):
                story.append(Paragraph("✅ Action Items", self.styles['Heading2']))
                for i, item in enumerate(summary['action_items'], 1):
                    story.append(Paragraph(f"{i}. {item}", info_style))
                story.append(Spacer(1, 20))
            
            # Key Decisions
            if summary.get('key_decisions'):
                story.append(Paragraph("🔑 Key Decisions", self.styles['Heading2']))
                for i, decision in enumerate(summary['key_decisions'], 1):
                    story.append(Paragraph(f"{i}. {decision}", info_style))
                story.append(Spacer(1, 20))
        
        # Next Meeting Agenda
        if export_data.get('next_agenda'):
            story.append(Paragraph("📅 Suggested Next Meeting Agenda", self.styles['Heading2']))
            for i, item in enumerate(export_data['next_agenda'], 1):
                story.append(Paragraph(f"{i}. {item}", info_style))
            story.append(Spacer(1, 20))
        
        # Topics
        topics = export_data.get('topics')
        if topics:
            story.append(Paragraph("🎯 Meeting Topics", self.styles['Heading2']))
            for topic in topics:
                story.append(Paragraph(f"<b>{topic['timestamp']} - {topic['title']}</b>", self.styles['Heading3']))
                story.append(Paragraph(f"Duration: {topic['duration']}", info_style))
                story.append(Paragraph(topic['summary'], info_style))
                story.append(Spacer(1, 15))
        
        # Full Transcript
        story.append(PageBreak())
        story.append(Paragraph("📝 Complete Meeting Transcript", self.styles['Heading2']))
        story.append(Spacer(1, 20))
        
        transcript = export_data.get('transcript', '')
        transcript_chunks = [transcript[i:i+2000] for i in range(0, len(transcript), 2000)]
        for chunk in transcript_chunks:
            story.append(Paragraph(chunk, info_style))
        
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
    
    def export_risk_report(self, export_data: dict):
        """Export focused risk analysis report"""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch)
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'RiskTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,
            textColor='red'
        )
        
        story.append(Paragraph("🚨 Meeting Risk Analysis Report", title_style))
        story.append(Spacer(1, 20))
        
        # Report metadata
        info_style = self.styles['Normal']
        story.append(Paragraph(f"<b>Generated on:</b> {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", info_style))
        story.append(Spacer(1, 20))
        
        risks = export_data.get('risks', {})
        
        # Executive Risk Summary
        story.append(Paragraph("📊 Risk Summary", self.styles['Heading2']))
        
        total_risks = risks.get('total_risks', 0)
        urgency_score = risks.get('urgency_score', 0)
        
        # Determine risk level
        if urgency_score >= 70 or total_risks >= 8:
            risk_level = "🔴 HIGH RISK"
            risk_color = 'red'
        elif urgency_score >= 40 or total_risks >= 4:
            risk_level = "🟡 MEDIUM RISK"
            risk_color = 'orange'
        else:
            risk_level = "🟢 LOW RISK"
            risk_color = 'green'
        
        risk_style = ParagraphStyle(
            'RiskLevel',
            parent=self.styles['Normal'],
            fontSize=16,
            textColor=risk_color,
            spaceAfter=20
        )
        
        story.append(Paragraph(f"<b>Overall Risk Level:</b> {risk_level}", risk_style))
        story.append(Paragraph(f"<b>Total Risk Items:</b> {total_risks}", info_style))
        story.append(Paragraph(f"<b>Urgency Score:</b> {urgency_score}/100", info_style))
        story.append(Spacer(1, 30))
        
        # Detailed Risk Categories
        risk_categories = [
            ('⏰ Critical Deadlines', risks.get('deadlines', []), 'red'),
            ('😠 Customer Issues', risks.get('customer_issues', []), 'red'),
            ('💰 Budget Risks', risks.get('budget_risks', []), 'orange'),
            ('⚖️ Legal Concerns', risks.get('legal_concerns', []), 'blue')
        ]
        
        for category_name, items, color in risk_categories:
            if items:
                category_style = ParagraphStyle(
                    f'Category_{color}',
                    parent=self.styles['Heading2'],
                    textColor=color
                )
                
                story.append(Paragraph(f"{category_name} ({len(items)} items)", category_style))
                
                for i, item in enumerate(items, 1):
                    story.append(Paragraph(f"{i}. {item}", info_style))
                
                story.append(Spacer(1, 20))
        
        # Recommendations
        story.append(Paragraph("💡 Recommended Actions", self.styles['Heading2']))
        
        if urgency_score >= 70:
            story.append(Paragraph("🚨 <b>IMMEDIATE ACTION REQUIRED</b>", risk_style))
            story.append(Paragraph("• Schedule emergency follow-up within 24 hours", info_style))
            story.append(Paragraph("• Notify all stakeholders immediately", info_style))
            story.append(Paragraph("• Create detailed action plan with specific deadlines", info_style))
        
        elif urgency_score >= 40:
            story.append(Paragraph("⚠️ <b>PROMPT ACTION NEEDED</b>", info_style))
            story.append(Paragraph("• Schedule follow-up within 48-72 hours", info_style))
            story.append(Paragraph("• Assign specific owners to each risk item", info_style))
            story.append(Paragraph("• Set calendar reminders for critical dates", info_style))
        
        else:
            story.append(Paragraph("📋 <b>MONITOR AND DOCUMENT</b>", info_style))
            story.append(Paragraph("• Review at next regular meeting", info_style))
            story.append(Paragraph("• Document all items for future reference", info_style))
            story.append(Paragraph("• Continue monitoring for any escalation", info_style))
        
        doc.build(story)
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data